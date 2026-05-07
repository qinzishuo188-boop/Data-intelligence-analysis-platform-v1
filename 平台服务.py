from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import uuid
from email.parser import BytesParser
from email.policy import default
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from io import StringIO
from math import cos, pi, sin
from pathlib import Path
from typing import Any
from urllib.parse import quote, unquote, urlparse

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from analysis_engine import build_analysis_output, choose_preferred_category_column, clean_table_enhanced, recommend_charts_enhanced
from ai_table_extractor import extract_structured_table_with_ai, get_ai_status
from text_recognition import robust_text_to_table

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None


ROOT_DIR = Path(__file__).resolve().parent
WEB_DIR = ROOT_DIR / "web"
UPLOAD_DIR = ROOT_DIR / "uploads"
CACHE_DIR = ROOT_DIR / "cache"
EXPORT_DIR = ROOT_DIR / "exports"
HOST = "0.0.0.0"
PORT = 8866
PREVIEW_ROWS = 50
CHINESE_FONT = Path(r"C:\Windows\Fonts\msyh.ttc")

LEGACY_DIRS = {
    "网页资源": WEB_DIR,
    "上传文件": UPLOAD_DIR,
    "数据缓存": CACHE_DIR,
    "导出结果": EXPORT_DIR,
}

METRIC_ALIASES = {
    "批发价": "批发价",
    "采购价": "采购价",
    "收盘价": "收盘价",
    "单价": "单价",
    "均价": "均价",
    "日均采购量": "日均采购量",
    "日均采购": "日均采购量",
    "采购量": "采购量",
    "月销量": "月销量",
    "销量": "销量",
    "同比增长": "同比",
    "同比": "同比",
    "环比": "环比",
    "增长": "增长",
    "上涨": "增长",
    "下降": "增长",
    "微降": "增长",
    "占比": "占比",
    "补贴金额": "补贴金额",
    "好评率": "好评率",
    "月收入": "月收入",
    "平均销售额": "平均销售额",
    "平均月销量": "平均月销量",
    "平均库存周转天数": "平均库存周转天数",
}

METRIC_PATTERN = re.compile(
    r"(?P<label>批发价|采购价|收盘价|单价|均价|日均采购量|日均采购|采购量|月销量|销量|同比增长|同比|环比|增长|上涨|下降|微降|占比|补贴金额|好评率|月收入|平均销售额|平均月销量|平均库存周转天数)"
    r"(?:为|是|达|约为|约|较上周|较上月|较去年同期|较前一交易日|截至|增长|上涨|下降|微降)?"
    r"\s*(?P<value>-?\d+(?:\.\d+)?)\s*"
    r"(?P<unit>%|个百分点|元/斤|元/升|元/瓶|元/提|元/件|元/天|元|万元|万|亿|斤|升|瓶|提|件|天|岁)?"
)
GENERIC_PERCENT_PATTERN = re.compile(
    r"(?P<item>[^，；。:\n]{2,40}?)(?:增速最快|增长最快|增长显著|表现突出|增速领先|增长)?"
    r"(?:，\d{4}年)?(?:同比|环比)?增长(?P<value>-?\d+(?:\.\d+)?)%"
)
CATEGORY_PATTERN = re.compile(r"((?:水果|粮油|日用品|农产品|3C数码|家电|服饰|家居|平台|渠道|行业)[^,，;；。]{0,8}?类)")
SUPPLIER_PATTERN = re.compile(r"(?P<category>[^,，;；。]{1,14}?类)由(?P<supplier>[^,，;；。]{2,28}?供应商)")
ATTR_PATTERNS = {
    "产地": re.compile(r"产地(?:为)?(?P<value>[^,，;；。]+)"),
    "供应商": re.compile(r"由(?P<value>[^,，;；。]{2,28}?供应商)"),
}
NUMERIC_HINTS = {"价", "量", "率", "金额", "收入", "增长", "同比", "环比", "占比", "销量", "周转"}


def ensure_dirs() -> None:
    for legacy_name, target in LEGACY_DIRS.items():
        legacy_dir = ROOT_DIR / legacy_name
        if legacy_dir.exists() and not target.exists():
            legacy_dir.rename(target)

    for folder in (WEB_DIR, UPLOAD_DIR, CACHE_DIR, EXPORT_DIR):
        folder.mkdir(parents=True, exist_ok=True)


def json_response(payload: dict[str, Any], status_code: int = 200) -> tuple[bytes, int]:
    def sanitize(value: Any) -> Any:
        if isinstance(value, dict):
            return {key: sanitize(item) for key, item in value.items()}
        if isinstance(value, list):
            return [sanitize(item) for item in value]
        if isinstance(value, tuple):
            return [sanitize(item) for item in value]
        if pd.isna(value):
            return None
        return value

    return json.dumps(sanitize(payload), ensure_ascii=False).encode("utf-8"), status_code


def find_node_exe() -> str | None:
    env_path = os.environ.get("NODE_EXE")
    if env_path and Path(env_path).exists():
        return env_path

    candidates = [
        ROOT_DIR / "runtime" / "node" / "bin" / "node.exe",
        ROOT_DIR / "runtime" / "node" / "node.exe",
        Path(r"C:\Program Files\nodejs\node.exe"),
        Path(r"C:\Program Files (x86)\nodejs\node.exe"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return shutil.which("node")


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if CHINESE_FONT.exists():
        return ImageFont.truetype(str(CHINESE_FONT), size=size)
    return ImageFont.load_default()


def read_csv(path: Path) -> pd.DataFrame:
    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb18030"):
        try:
            return pd.read_csv(path, encoding=encoding)
        except Exception:
            continue
    raise ValueError("CSV 文件读取失败，请检查编码或内容格式。")


def json_to_table(data: Any) -> pd.DataFrame:
    if isinstance(data, list) and data and all(isinstance(item, dict) for item in data):
        return pd.DataFrame(data)
    if isinstance(data, dict):
        if all(isinstance(value, list) for value in data.values()):
            return pd.DataFrame(data)
        return pd.DataFrame([{"字段": key, "值": value} for key, value in data.items()])
    raise ValueError("JSON 内容暂不支持自动转换为表格。")


def extract_word_text(path: Path) -> str:
    if Document is None:
        raise RuntimeError("当前环境未安装 python-docx，无法解析 Word 文件。")
    doc = Document(path)
    lines = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_pdf_text(path: Path) -> str:
    if PdfReader is None:
        raise RuntimeError("当前环境未安装 pypdf，无法解析 PDF 文件。")
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def normalize_text(text: str) -> str:
    mapping = str.maketrans({
        "（": "(",
        "）": ")",
        "【": "[",
        "】": "]",
        "％": "%",
        "：": ":",
        "，": ",",
        "；": ";",
        "　": " ",
        "\ufeff": "",
        "\u200b": "",
        "\u200c": "",
        "\u200d": "",
    })
    text = (text or "").translate(mapping)
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def clean_label(value: str) -> str:
    text = str(value or "").strip(" ,;:()[]")
    text = re.sub(r"^(其中|此外|另外|同时|当前|本次|围绕|具体数据如下|从细分品类来看)", "", text)
    return text.strip() or "文本"


def clean_item_name(value: str) -> str:
    text = str(value or "").strip(" ,;:()[]")
    text = re.sub(r"^(其中|此外|另外|同时|当前|本次|围绕|具体数据如下|从细分品类来看|粮油类中|日用品类中)", "", text)
    text = re.sub(r"(增速最快|增长最快|增长显著|表现突出|增速领先)$", "", text)
    text = re.sub(r"[，,]?\d{4}年.*$", "", text)
    text = re.split(r"[:：]", text)[-1]
    return text.strip() or "项目"


def sanitize_source_name(name: str) -> str:
    stem = Path(name).stem or str(name)
    stem = re.sub(r"[_-]+", " ", stem).strip()
    return stem[:48] or "数据集"


def split_segments(text: str) -> list[str]:
    normalized = normalize_text(text)
    if not normalized:
        return []
    return [segment.strip() for segment in re.split(r"[;\n。]", normalized) if segment.strip()]


def parse_delimited_table(text: str) -> pd.DataFrame | None:
    lines = [line.strip() for line in normalize_text(text).splitlines() if line.strip()]
    for separator in (",", "\t", "|", ";"):
        matched = [line for line in lines if separator in line]
        if len(matched) < 2:
            continue
        try:
            table = pd.read_csv(StringIO("\n".join(matched)), sep=re.escape(separator), engine="python")
            if table.shape[1] >= 2 and len(table) >= 1:
                return table
        except Exception:
            continue
    return None


def parse_key_value_lines(text: str) -> pd.DataFrame | None:
    records: list[dict[str, Any]] = []
    for line in normalize_text(text).splitlines():
        line = line.strip()
        if not line:
            continue
        match = re.match(r"^(.+?)[:：]\s*(-?\d+(?:\.\d+)?)\s*$", line)
        if match:
            records.append({"项目": clean_item_name(match.group(1)), "值": float(match.group(2))})
    return pd.DataFrame(records) if len(records) >= 2 else None


def parse_supplier_map(text: str) -> dict[str, str]:
    supplier_map: dict[str, str] = {}
    for match in SUPPLIER_PATTERN.finditer(normalize_text(text)):
        supplier_map[clean_item_name(match.group("category"))] = clean_item_name(match.group("supplier"))
    return supplier_map


def extract_attributes(segment: str) -> dict[str, Any]:
    attrs: dict[str, Any] = {}
    bracket_match = re.search(r"\(([^)]+)\)", segment)
    if bracket_match:
        attrs["规格"] = clean_item_name(bracket_match.group(1))
    for pattern in (r"(果径[^,，;；。]+)", r"(\d+\s*ml瓶装)", r"(\d+\s*卷装)", r"(纸袋包装)"):
        match = re.search(pattern, segment, flags=re.IGNORECASE)
        if match and "规格" not in attrs:
            attrs["规格"] = clean_item_name(match.group(1))
    for label, pattern in ATTR_PATTERNS.items():
        match = pattern.search(segment)
        if match:
            attrs[label] = clean_item_name(match.group("value"))
    return attrs


def extract_metrics(segment: str) -> tuple[dict[str, Any], int | None]:
    values: dict[str, Any] = {}
    first_index: int | None = None
    for match in METRIC_PATTERN.finditer(segment):
        if first_index is None:
            first_index = match.start()
        raw_label = match.group("label")
        label = METRIC_ALIASES.get(raw_label, raw_label)
        unit = match.group("unit") or ""
        value = float(match.group("value"))
        matched_text = match.group(0)
        if raw_label in {"下降", "微降"} or any(word in matched_text for word in ("下降", "下跌", "微降")):
            value = -abs(value)
        column_name = f"{label}({unit})" if unit and unit not in {"%", "个百分点"} else label
        values[column_name] = value
    return values, first_index


def infer_category(segment: str, current_category: str) -> str:
    matches = CATEGORY_PATTERN.findall(segment)
    if matches:
        return clean_item_name(matches[-1])
    return current_category


def choose_numeric_columns(table: pd.DataFrame) -> list[str]:
    numeric_columns = [column for column in table.columns if pd.api.types.is_numeric_dtype(table[column])]
    def score(column: str) -> tuple[int, int]:
        non_null = int(pd.to_numeric(table[column], errors="coerce").notna().sum())
        priority = 0
        if any(keyword in column for keyword in ("采购量", "销量", "销售额", "同比", "增长", "占比", "好评率")):
            priority = 1
        return (non_null, priority)
    return sorted(numeric_columns, key=score, reverse=True)


def infer_item_name(segment: str, first_metric_index: int | None) -> str:
    prefix = segment if first_metric_index is None else segment[:first_metric_index]
    prefix = re.split(r"[:：]", prefix)[-1]
    prefix = re.sub(r"\(([^)]*)\)", "", prefix)
    prefix = re.sub(r"^(以|由|围绕)", "", prefix)
    prefix = re.sub(r"(作为核心|为核心|涵盖.*)$", "", prefix)
    prefix = re.sub(r"(增速最快|增长最快|增长显著|表现突出|增速领先)$", "", prefix)
    prefix = re.sub(r"[，,]?\d{4}年.*$", "", prefix)
    parts = [part.strip() for part in re.split(r"[，,]", prefix) if part.strip()]
    if parts:
        prefix = parts[-1]
    return clean_item_name(prefix)


def parse_narrative_records(text: str) -> pd.DataFrame | None:
    normalized = normalize_text(text)
    segments = split_segments(normalized)
    if len(segments) < 2:
        return None

    supplier_map = parse_supplier_map(normalized)
    records: list[dict[str, Any]] = []
    current_category = ""

    for segment in segments:
        current_category = infer_category(segment, current_category)
        metrics, first_metric_index = extract_metrics(segment)

        if not metrics:
            generic_match = GENERIC_PERCENT_PATTERN.search(segment)
            if generic_match:
                item_name = clean_item_name(generic_match.group("item"))
                category = infer_category(item_name, current_category) or item_name
                records.append({
                    "分类": category,
                    "项目": item_name,
                    "同比": float(generic_match.group("value")),
                    "原文片段": segment,
                })
            continue

        item_name = infer_item_name(segment, first_metric_index)
        if len(item_name) > 42:
            continue

        record: dict[str, Any] = {
            "分类": current_category or "未分类",
            "项目": item_name,
            "原文片段": segment,
        }
        record.update(extract_attributes(segment))
        record.update(metrics)

        category = record["分类"]
        if category in supplier_map and "供应商" not in record:
            record["供应商"] = supplier_map[category]
        elif "供应商" not in record:
            segment_category = infer_category(segment, "")
            if segment_category in supplier_map:
                record["供应商"] = supplier_map[segment_category]

        records.append(record)

    if len(records) < 2:
        return None
    return pd.DataFrame(records)


def parse_text_fragments(text: str) -> pd.DataFrame:
    return pd.DataFrame({
        "片段内容": split_segments(text),
    })


def text_to_table(text: str) -> pd.DataFrame:
    normalized = normalize_text(text)
    for parser in (parse_delimited_table, parse_key_value_lines, parse_narrative_records):
        table = parser(normalized)
        if table is not None and not table.empty:
            return table
    robust_table = robust_text_to_table(normalized)
    if robust_table is not None and not robust_table.empty and "片段内容" not in robust_table.columns:
        return robust_table
    return parse_text_fragments(normalized)


def file_to_table(path: Path) -> tuple[pd.DataFrame, str, str]:
    suffix = path.suffix.lower()
    raw_text = ""
    source_name = sanitize_source_name(path.name)
    if suffix == ".csv":
        return read_csv(path), source_name, raw_text
    if suffix in {".xlsx", ".xls"}:
        workbook = pd.ExcelFile(path)
        sheet_name = workbook.sheet_names[0]
        table = pd.read_excel(path, sheet_name=sheet_name)
        return table, source_name, raw_text
    if suffix == ".json":
        return json_to_table(json.loads(path.read_text(encoding="utf-8"))), source_name, raw_text
    if suffix == ".txt":
        raw_text = path.read_text(encoding="utf-8", errors="ignore")
        return text_to_table(raw_text), source_name, raw_text
    if suffix == ".docx":
        raw_text = extract_word_text(path)
        return text_to_table(raw_text), source_name, raw_text
    if suffix == ".pdf":
        raw_text = extract_pdf_text(path)
        return text_to_table(raw_text), source_name, raw_text
    raise ValueError(f"暂不支持该文件类型：{suffix}")


def maybe_extract_with_ai(raw_text: str, source_name: str) -> tuple[pd.DataFrame | None, dict[str, Any]]:
    text = (raw_text or "").strip()
    if not text:
        return None, {"mode": "rule", "used": False}

    try:
        table, meta = extract_structured_table_with_ai(text, source_name)
        meta["used"] = True
        return table, meta
    except Exception as exc:
        return None, {
            "mode": "rule",
            "used": False,
            "fallbackReason": str(exc),
        }


def parse_text_input(text: str, source_name: str) -> tuple[pd.DataFrame, dict[str, Any]]:
    stripped = (text or "").strip()
    if not stripped:
        raise ValueError("璇疯緭鍏ユ枃鏈唴瀹瑰悗鍐嶇敓鎴愬浘琛ㄣ€?")

    try:
        return json_to_table(json.loads(stripped)), {
            "mode": "json",
            "used": False,
            "parser": "json",
        }
    except Exception:
        pass

    ai_table, extraction_meta = maybe_extract_with_ai(stripped, source_name)
    if ai_table is not None:
        return ai_table, extraction_meta

    rule_table = text_to_table(stripped)
    extraction_meta["parser"] = "rule"
    return rule_table, extraction_meta


def parse_text_input_by_mode(text: str, source_name: str, mode: str = "auto") -> tuple[pd.DataFrame, dict[str, Any]]:
    stripped = (text or "").strip()
    if not stripped:
        raise ValueError("璇疯緭鍏ユ枃鏈唴瀹瑰悗鍐嶇敓鎴愬浘琛ㄣ€?")

    if mode == "rule":
        try:
            return json_to_table(json.loads(stripped)), {
                "mode": "json",
                "used": False,
                "parser": "json",
            }
        except Exception:
            table = text_to_table(stripped)
            return table, {
                "mode": "rule",
                "used": False,
                "parser": "rule",
            }

    if mode == "ai":
        ai_table, extraction_meta = maybe_extract_with_ai(stripped, source_name)
        if ai_table is None:
            raise RuntimeError(extraction_meta.get("fallbackReason") or "AI 提取失败")
        return ai_table, extraction_meta

    return parse_text_input(stripped, source_name)


def clean_table(table: pd.DataFrame) -> pd.DataFrame:
    return clean_table_enhanced(table)
    cleaned = table.copy()
    cleaned.columns = [
        str(column).strip() if str(column).strip() else f"字段{index + 1}"
        for index, column in enumerate(cleaned.columns)
    ]
    cleaned = cleaned.dropna(how="all").reset_index(drop=True)
    for column in cleaned.columns:
        if cleaned[column].dtype == object:
            converted = pd.to_numeric(cleaned[column], errors="coerce")
            if converted.notna().sum() >= max(2, len(cleaned) // 2):
                cleaned[column] = converted
    return cleaned


def infer_story_mode(table: pd.DataFrame) -> str:
    numeric_columns = [column for column in table.columns if pd.api.types.is_numeric_dtype(table[column])]
    return "结构化指标" if numeric_columns else "文本抽取"


def recommend_charts(table: pd.DataFrame, numeric_columns: list[str], category_columns: list[str]) -> list[dict[str, Any]]:
    return recommend_charts_enhanced(table, numeric_columns, category_columns)
    suggestions: list[dict[str, Any]] = []
    if numeric_columns and category_columns:
        x_field = category_columns[0]
        primary = numeric_columns[: min(3, len(numeric_columns))]
        suggestions.extend([
            {"type": "horizontalBar", "x": x_field, "y": primary[:1], "title": f"{x_field} 横向对比图"},
            {"type": "bar", "x": x_field, "y": primary[:1], "title": f"{x_field} 柱状图"},
            {"type": "line", "x": x_field, "y": primary[:1], "title": f"{x_field} 趋势折线图"},
            {"type": "pie", "x": x_field, "y": primary[:1], "title": f"{x_field} 占比饼图"},
            {"type": "rosePie", "x": x_field, "y": primary[:1], "title": f"{x_field} 玫瑰图"},
            {"type": "funnel", "x": x_field, "y": primary[:1], "title": f"{x_field} 漏斗图"},
        ])
        if len(primary) >= 2:
            suggestions.extend([
                {"type": "stackedBar", "x": x_field, "y": primary, "title": f"{x_field} 堆叠柱状图"},
                {"type": "radar", "x": x_field, "y": primary[: min(5, len(primary))], "title": f"{x_field} 雷达图"},
            ])
    elif len(numeric_columns) >= 2:
        suggestions.append({"type": "scatter", "x": numeric_columns[0], "y": [numeric_columns[1]], "title": "双指标散点图"})
        suggestions.append({"type": "bar", "x": "", "y": [numeric_columns[0]], "title": f"{numeric_columns[0]} 排名图"})
    elif len(numeric_columns) == 1:
        suggestions.extend([
            {"type": "bar", "x": category_columns[0] if category_columns else "", "y": [numeric_columns[0]], "title": f"{numeric_columns[0]} 柱状图"},
            {"type": "line", "x": category_columns[0] if category_columns else "", "y": [numeric_columns[0]], "title": f"{numeric_columns[0]} 折线图"},
        ])
    return suggestions[:8]


def format_number(value: float) -> str:
    if abs(value) >= 100:
        return f"{value:,.0f}"
    return f"{value:,.2f}"


def top_record_text(table: pd.DataFrame, category_column: str, metric_column: str) -> str | None:
    series = pd.to_numeric(table[metric_column], errors="coerce")
    valid = table.loc[series.notna()].copy()
    if valid.empty:
        return None
    max_row = valid.loc[series[series.notna()].idxmax()]
    return f"{metric_column} 最高的是 {max_row.get(category_column, max_row.iloc[0])}，达到 {format_number(float(max_row[metric_column]))}。"


def build_summary(table: pd.DataFrame, numeric_columns: list[str], category_columns: list[str], raw_text: str) -> dict[str, Any]:
    cards = [
        {"label": "记录数", "value": str(len(table))},
        {"label": "字段数", "value": str(len(table.columns))},
        {"label": "数值指标", "value": str(len(numeric_columns))},
        {"label": "识别模式", "value": infer_story_mode(table)},
    ]

    insights: list[str] = []
    for column in numeric_columns[:3]:
        series = pd.to_numeric(table[column], errors="coerce").dropna()
        if not series.empty:
            insights.append(
                f"{column} 的均值为 {format_number(float(series.mean()))}，最大值为 {format_number(float(series.max()))}。"
            )

    if category_columns and numeric_columns:
        text = top_record_text(table, category_columns[0], numeric_columns[0])
        if text:
            insights.append(text)

    if "供应商" in table.columns:
        supplier_count = table["供应商"].fillna("").astype(str).str.strip().replace("", pd.NA).dropna().nunique()
        if supplier_count:
            insights.append(f"当前数据共识别出 {supplier_count} 个供应商主体，便于直接形成采购分工汇报。")

    if raw_text.strip():
        insights.append("原始长段文字已同步拆分为结构化表格，适合继续加工成经营汇报和展示材料。")

    story: list[str] = []
    if category_columns and numeric_columns:
        story.append(f"建议以“{category_columns[0]} + {numeric_columns[0]}”作为主图，先展示整体差异。")
    if len(numeric_columns) >= 2 and category_columns:
        story.append("第二张图可切换为堆叠柱状图或雷达图，用于补充多指标对比。")
    story.append("结构化表格已同步生成，适合在 PPT 中直接作为核心数据页展示。")
    story.append("导出稿件已去掉原始文件名和工作表信息，展示口径会更干净。")

    return {"cards": cards, "insights": insights[:6], "story": story[:4]}


def build_dataset(
    table: pd.DataFrame,
    source_name: str,
    raw_text: str = "",
    dataset_id: str | None = None,
    source_kind: str = "file",
    extraction_meta: dict[str, Any] | None = None,
) -> dict[str, Any]:
    table = clean_table(table)
    numeric_columns = choose_numeric_columns(table)
    category_columns = [column for column in table.columns if column not in numeric_columns]
    preferred_category = choose_preferred_category_column(table, category_columns)
    if preferred_category:
        category_columns = [preferred_category] + [column for column in category_columns if column != preferred_category]
    preview = table.head(PREVIEW_ROWS)
    analysis = build_analysis_output(table, source_name, numeric_columns, category_columns, raw_text)
    return {
        "datasetId": dataset_id,
        "sourceName": source_name,
        "sourceKind": source_kind,
        "rowCount": int(len(table)),
        "columnCount": int(len(table.columns)),
        "columns": [{"name": column, "numeric": column in numeric_columns} for column in table.columns],
        "numericColumns": numeric_columns,
        "categoryColumns": category_columns,
        "rows": preview.where(pd.notnull(preview), None).to_dict(orient="records"),
        "suggestions": analysis["charts"],
        "rawTextPreview": raw_text[:2400],
        "title": analysis["title"],
        "summary": analysis["summary"],
        "tables": analysis["tables"],
        "charts": analysis["charts"],
        "insights": analysis["insights"],
        "recognizedFields": analysis["recognizedFields"],
        "analysis": analysis,
        "extraction": extraction_meta or {"mode": "rule", "used": False},
    }
    table = clean_table(table)
    numeric_columns = choose_numeric_columns(table)
    category_columns = [column for column in table.columns if column not in numeric_columns]
    preview = table.head(PREVIEW_ROWS)
    return {
        "datasetId": dataset_id,
        "sourceName": source_name,
        "sourceKind": source_kind,
        "rowCount": int(len(table)),
        "columnCount": int(len(table.columns)),
        "columns": [{"name": column, "numeric": column in numeric_columns} for column in table.columns],
        "numericColumns": numeric_columns,
        "categoryColumns": category_columns,
        "rows": preview.where(pd.notnull(preview), None).to_dict(orient="records"),
        "suggestions": recommend_charts(table, numeric_columns, category_columns),
        "rawTextPreview": raw_text[:2400],
        "summary": build_summary(table, numeric_columns, category_columns, raw_text),
    }


def save_dataset(dataset_id: str, table: pd.DataFrame, dataset: dict[str, Any]) -> None:
    folder = CACHE_DIR / dataset_id
    folder.mkdir(parents=True, exist_ok=True)
    table.to_csv(folder / "data.csv", index=False, encoding="utf-8-sig")
    (folder / "meta.json").write_text(json.dumps(dataset, ensure_ascii=False, indent=2), encoding="utf-8")


def load_dataset(dataset_id: str) -> tuple[pd.DataFrame, dict[str, Any]]:
    folder = CACHE_DIR / dataset_id
    data_file = folder / "data.csv"
    meta_file = folder / "meta.json"
    if not data_file.exists() or not meta_file.exists():
        raise FileNotFoundError("没有找到对应的数据集，请重新上传。")
    return clean_table(pd.read_csv(data_file)), json.loads(meta_file.read_text(encoding="utf-8"))


def parse_upload(content_type: str, request_body: bytes) -> tuple[str, bytes]:
    message = BytesParser(policy=default).parsebytes(
        f"Content-Type: {content_type}\r\nMIME-Version: 1.0\r\n\r\n".encode("utf-8") + request_body
    )
    if not message.is_multipart():
        raise ValueError("上传请求格式不正确。")
    for part in message.iter_parts():
        if part.get_content_disposition() != "form-data":
            continue
        field_name = part.get_param("name", header="content-disposition")
        filename = part.get_filename()
        if field_name == "file" and filename:
            return Path(filename).name, part.get_payload(decode=True) or b""
    raise ValueError("请选择要上传的文件。")


def choose_category_field(table: pd.DataFrame, config: dict[str, Any]) -> str:
    x_field = config.get("xField") or ""
    if x_field and x_field in table.columns:
        return x_field
    for candidate in ("项目", "分类", "名称", "标题", "字段"):
        if candidate in table.columns:
            return candidate
    return str(table.columns[0])


def choose_numeric_fields(table: pd.DataFrame, config: dict[str, Any]) -> list[str]:
    requested = [field for field in config.get("yFields", []) if field in table.columns]
    if requested:
      return requested
    numeric_columns = [column for column in table.columns if pd.api.types.is_numeric_dtype(table[column])]
    return numeric_columns[:3]


def text_width(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> int:
    left, _, right, _ = draw.textbbox((0, 0), text, font=font)
    return right - left


def fit_label(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> str:
    label = str(text)
    if text_width(draw, label, font) <= max_width:
        return label
    ellipsis = "..."
    while label and text_width(draw, f"{label}{ellipsis}", font) > max_width:
        label = label[:-1]
    return f"{label}{ellipsis}" if label else ellipsis


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str, width: int) -> None:
    draw.rounded_rectangle((36, 24, width - 36, 116), radius=22, outline="#d6e3f1", width=2)


def draw_axis(draw: ImageDraw.ImageDraw, left: int, top: int, right: int, bottom: int) -> None:
    draw.line((left, top, left, bottom), fill="#86a3c4", width=2)
    draw.line((left, bottom, right, bottom), fill="#86a3c4", width=2)


def render_bar_like(draw: ImageDraw.ImageDraw, width: int, height: int, categories: list[str], series_map: dict[str, list[float]], config: dict[str, Any], horizontal: bool = False) -> None:
    chart_left = 260 if horizontal else 92
    chart_top, chart_right, chart_bottom = 170, width - 56, height - 130
    draw.rounded_rectangle((56, 136, width - 44, height - 58), radius=24, fill="#ffffff", outline="#d6e3f1", width=2)
    draw_axis(draw, chart_left, chart_top, chart_right, chart_bottom)

    all_values = [value for values in series_map.values() for value in values] or [0]
    max_value = max(max(all_values), 1)
    min_value = min(min(all_values), 0)
    if min_value > 0:
        min_value = 0
    font = load_font(18)

    for tick in range(5):
        y = chart_bottom - (chart_bottom - chart_top) * tick / 4
        label_value = min_value + (max_value - min_value) * tick / 4
        if not horizontal:
            draw.line((chart_left, y, chart_right, y), fill="#eef4fb", width=1)
            draw.text((28, y - 10), format_number(float(label_value)), fill="#6b88a8", font=font)
        else:
            x = chart_left + (chart_right - chart_left) * tick / 4
            draw.line((x, chart_top, x, chart_bottom), fill="#eef4fb", width=1)
            draw.text((x - 18, chart_bottom + 16), format_number(float(label_value)), fill="#6b88a8", font=font)

    colors = ["#1f6fbf", "#138f84", "#c47a18", "#8e5bb5"]
    series_names = list(series_map.keys())
    if horizontal:
        row_height = (chart_bottom - chart_top) / max(len(categories), 1)
        bar_height = max(row_height / max(len(series_names), 1) - 10, 10)
        for category_index, category in enumerate(categories):
            y_base = chart_top + category_index * row_height + 10
            label = fit_label(draw, category, font, chart_left - 62)
            label_x = chart_left - 20 - text_width(draw, label, font)
            draw.text((label_x, y_base), label, fill="#385473", font=font)
            for series_index, series_name in enumerate(series_names):
                value = series_map[series_name][category_index]
                x1 = chart_left + (value - min_value) / (max_value - min_value or 1) * (chart_right - chart_left)
                y0 = y_base + series_index * (bar_height + 6)
                y1 = y0 + bar_height
                draw.rounded_rectangle((chart_left, y0, x1, y1), radius=6, fill=colors[series_index % len(colors)])
    else:
        step = (chart_right - chart_left) / max(len(categories), 1)
        bar_width = max((step - 20) / max(len(series_names), 1), 10)
        for series_index, series_name in enumerate(series_names):
            for category_index, value in enumerate(series_map[series_name]):
                x0 = chart_left + category_index * step + 10 + series_index * bar_width
                x1 = x0 + bar_width - 6
                y0 = chart_bottom - (value - min_value) / (max_value - min_value or 1) * (chart_bottom - chart_top)
                draw.rounded_rectangle((x0, y0, x1, chart_bottom), radius=6, fill=colors[series_index % len(colors)])
            draw.rounded_rectangle((chart_left + series_index * 180, 146, chart_left + series_index * 180 + 22, 168), radius=5, fill=colors[series_index % len(colors)])
            draw.text((chart_left + series_index * 180 + 32, 144), series_name[:16], fill="#294a71", font=font)
        for category_index, category in enumerate(categories):
            label = category[:10] + ("..." if len(category) > 10 else "")
            x = chart_left + category_index * step + step / 2 - 24
            draw.text((x, chart_bottom + 16), label, fill="#385473", font=font)


def render_line_chart(draw: ImageDraw.ImageDraw, width: int, height: int, categories: list[str], series_map: dict[str, list[float]]) -> None:
    chart_left, chart_top, chart_right, chart_bottom = 92, 170, width - 56, height - 130
    draw.rounded_rectangle((56, 136, width - 44, height - 58), radius=24, fill="#ffffff", outline="#d6e3f1", width=2)
    draw_axis(draw, chart_left, chart_top, chart_right, chart_bottom)

    all_values = [value for values in series_map.values() for value in values] or [0]
    max_value = max(max(all_values), 1)
    min_value = min(min(all_values), 0)
    if min_value > 0:
        min_value = 0
    colors = ["#1f6fbf", "#138f84", "#c47a18", "#8e5bb5"]
    font = load_font(18)
    step = (chart_right - chart_left) / max(len(categories) - 1, 1)

    for tick in range(5):
        y = chart_bottom - (chart_bottom - chart_top) * tick / 4
        label_value = min_value + (max_value - min_value) * tick / 4
        draw.line((chart_left, y, chart_right, y), fill="#eef4fb", width=1)
        draw.text((28, y - 10), format_number(float(label_value)), fill="#6b88a8", font=font)

    for series_index, (series_name, values) in enumerate(series_map.items()):
        points: list[tuple[float, float]] = []
        for category_index, value in enumerate(values):
            x = chart_left + category_index * step
            y = chart_bottom - (value - min_value) / (max_value - min_value or 1) * (chart_bottom - chart_top)
            points.append((x, y))
        draw.line(points, fill=colors[series_index % len(colors)], width=4)
        for x, y in points:
            draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill=colors[series_index % len(colors)])
        draw.rounded_rectangle((chart_left + series_index * 180, 146, chart_left + series_index * 180 + 22, 168), radius=5, fill=colors[series_index % len(colors)])
        draw.text((chart_left + series_index * 180 + 32, 144), series_name[:16], fill="#294a71", font=font)

    for category_index, category in enumerate(categories):
        label = category[:10] + ("..." if len(category) > 10 else "")
        x = chart_left + category_index * step - 24
        draw.text((x, chart_bottom + 16), label, fill="#385473", font=font)


def render_pie_chart(draw: ImageDraw.ImageDraw, width: int, height: int, categories: list[str], values: list[float], rose: bool = False) -> None:
    draw.rounded_rectangle((56, 136, width - 44, height - 58), radius=24, fill="#ffffff", outline="#d6e3f1", width=2)
    box = (130, 180, 760, 810)
    total = sum(abs(value) for value in values) or 1
    colors = ["#67c2ff", "#2dd4bf", "#ffb65c", "#ff8f70", "#8f7cf7", "#5e97f6"]
    start = -pi / 2
    font = load_font(18)
    for index, value in enumerate(values):
        portion = abs(value) / total
        end = start + portion * 2 * pi
        radius = 250
        if rose:
            radius = 160 + int(120 * (abs(value) / (max(abs(v) for v in values) or 1)))
            center_x, center_y = 445, 495
            points = [(center_x, center_y)]
            for step in range(20):
                angle = start + (end - start) * step / 19
                points.append((center_x + radius * cos(angle), center_y + radius * sin(angle)))
            draw.polygon(points, fill=colors[index % len(colors)])
        else:
            draw.pieslice(box, start * 180 / pi, end * 180 / pi, fill=colors[index % len(colors)])
        start = end
    if not rose:
        draw.ellipse((285, 335, 605, 655), fill="#ffffff")

    legend_y = 196
    for index, category in enumerate(categories[:8]):
        draw.rounded_rectangle((820, legend_y, 842, legend_y + 22), radius=5, fill=colors[index % len(colors)])
        draw.text((854, legend_y - 2), f"{category[:14]}  {format_number(float(values[index]))}", fill="#294a71", font=font)
        legend_y += 38


def render_funnel(draw: ImageDraw.ImageDraw, width: int, height: int, categories: list[str], values: list[float]) -> None:
    draw.rounded_rectangle((56, 136, width - 44, height - 58), radius=24, fill="#ffffff", outline="#d6e3f1", width=2)
    max_value = max(values) if values else 1
    colors = ["#1f6fbf", "#2aa39b", "#ffb65c", "#ff8f70", "#8f7cf7"]
    font = load_font(20)
    top = 210
    height_step = 92
    center_x = width // 2 - 80
    for index, (category, value) in enumerate(zip(categories, values)):
        ratio = value / (max_value or 1)
        top_width = 520 * ratio + 120
        bottom_ratio = (values[index + 1] / (max_value or 1)) if index + 1 < len(values) else ratio * 0.6
        bottom_width = 520 * bottom_ratio + 120
        y0 = top + index * height_step
        y1 = y0 + height_step - 8
        points = [
            (center_x - top_width / 2, y0),
            (center_x + top_width / 2, y0),
            (center_x + bottom_width / 2, y1),
            (center_x - bottom_width / 2, y1),
        ]
        draw.polygon(points, fill=colors[index % len(colors)])
        draw.text((center_x - top_width / 2 + 18, y0 + 24), f"{category[:16]}  {format_number(float(value))}", fill="#ffffff", font=font)


def export_png(table: pd.DataFrame, config: dict[str, Any], output_dir: Path, title: str) -> Path:
    width, height = 1600, 980
    image = Image.new("RGB", (width, height), "#f7fbff")
    draw = ImageDraw.Draw(image)
    chart_type = config["type"]
    category_field = choose_category_field(table, config)
    numeric_fields = choose_numeric_fields(table, config)
    categories = [str(value) for value in table[category_field].fillna("").astype(str).tolist()[:12]] if category_field in table.columns else [f"第{i + 1}项" for i in range(min(len(table), 12))]
    series_map: dict[str, list[float]] = {}
    for field in numeric_fields[:4]:
        series_map[field] = [float(value or 0) for value in pd.to_numeric(table[field], errors="coerce").fillna(0).tolist()[: len(categories)]]

    draw_title(draw, title, "", width)

    if chart_type in {"pie", "rosePie"} and numeric_fields:
        render_pie_chart(draw, width, height, categories, series_map[numeric_fields[0]], rose=chart_type == "rosePie")
    elif chart_type == "funnel" and numeric_fields:
        render_funnel(draw, width, height, categories, series_map[numeric_fields[0]])
    elif chart_type in {"line", "area"} and series_map:
        render_line_chart(draw, width, height, categories, series_map)
    else:
        render_bar_like(draw, width, height, categories, series_map or {"值": [0]}, config, horizontal=chart_type == "horizontalBar")

    path = output_dir / "chart_report.png"
    image.save(path, compress_level=1, optimize=False)
    return path


def export_excel(table: pd.DataFrame, dataset: dict[str, Any], config: dict[str, Any], output_dir: Path) -> Path:
    path = output_dir / "chart_report.xlsx"
    overview_rows = [
        {"项目": "标题", "内容": dataset.get("title", "数据智能分析")},
        {"项目": "摘要", "内容": dataset.get("summary", "")},
        {"项目": "图表类型", "内容": config["type"]},
        {"项目": "分类字段", "内容": config.get("xField") or "自动索引"},
        {"项目": "数值字段", "内容": "、".join(config.get("yFields", []))},
    ]
    overview_rows.extend(
        {"项目": f"洞察{i + 1}", "内容": text}
        for i, text in enumerate(dataset.get("insights", []))
    )
    chart_rows = [
        {"图表标题": item.get("title", ""), "图表类型": item.get("type", ""), "推荐原因": item.get("reason", "")}
        for item in dataset.get("charts", [])
    ]
    field_rows = [
        {"字段": key, "识别结果": "、".join(values)}
        for key, values in (dataset.get("recognizedFields") or {}).items()
        if values
    ]

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        table.to_excel(writer, index=False, sheet_name="Data")
        pd.DataFrame(overview_rows).to_excel(writer, index=False, sheet_name="Overview")
        pd.DataFrame(chart_rows or [{"图表标题": "", "图表类型": "", "推荐原因": ""}]).to_excel(writer, index=False, sheet_name="Charts")
        pd.DataFrame(field_rows or [{"字段": "", "识别结果": ""}]).to_excel(writer, index=False, sheet_name="Fields")
    return path
    path = output_dir / "chart_report.xlsx"
    summary = dataset.get("summary", {})
    rows = [{"项目": card["label"], "内容": card["value"]} for card in summary.get("cards", [])]
    rows.extend({"项目": f"洞察{i + 1}", "内容": text} for i, text in enumerate(summary.get("insights", [])))
    rows.extend({"项目": f"表达{i + 1}", "内容": text} for i, text in enumerate(summary.get("story", [])))
    rows.extend([
        {"项目": "图表类型", "内容": config["type"]},
        {"项目": "分类字段", "内容": config.get("xField") or "自动索引"},
        {"项目": "数值字段", "内容": "、".join(config.get("yFields", []))},
    ])

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        table.to_excel(writer, index=False, sheet_name="Data")
        pd.DataFrame(rows).to_excel(writer, index=False, sheet_name="Summary")
    return path


def export_ppt(table: pd.DataFrame, dataset: dict[str, Any], image_path: Path, config: dict[str, Any], output_dir: Path) -> Path:
    node_exe = find_node_exe()
    if not node_exe:
        raise RuntimeError("未找到 Node.js，无法导出 PPT。请先安装 Node.js。")

    spec = {
        "title": dataset.get("title") or "数据智能分析汇报",
        "headline": dataset.get("title") or "数据智能分析汇报",
        "summaryText": dataset.get("summary", ""),
        "insights": dataset.get("insights", []),
        "chartIdeas": [item.get("title", "") for item in dataset.get("charts", [])[:5]],
        "chartImage": str(image_path),
        "pptxPath": str(output_dir / "chart_report.pptx"),
        "columns": list(table.columns),
        "rows": table.head(8).fillna("").astype(str).to_dict(orient="records"),
        "chartType": config["type"],
    }
    spec_path = output_dir / "ppt_spec.json"
    spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")

    script_path = ROOT_DIR / "export_ppt_report.mjs"
    result = subprocess.run(
        [str(node_exe), str(script_path), str(spec_path)],
        cwd=str(ROOT_DIR),
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or "PPT 导出失败")
    return output_dir / "chart_report.pptx"
    node_exe = find_node_exe()
    if not node_exe:
        raise RuntimeError("未找到 Node.js，无法导出 PPT。请先安装 Node.js。")

    spec = {
        "title": dataset.get("summary", {}).get("story", ["多元化图表汇报"])[0].replace("建议以“", "").replace("”作为主图，先展示整体差异。", "") or "多元化图表汇报",
        "headline": "多元化图表汇报",
        "summary": dataset.get("summary", {}),
        "chartImage": str(image_path),
        "pptxPath": str(output_dir / "chart_report.pptx"),
        "columns": list(table.columns),
        "rows": table.head(8).fillna("").astype(str).to_dict(orient="records"),
        "chartType": config["type"],
    }
    spec_path = output_dir / "ppt_spec.json"
    spec_path.write_text(json.dumps(spec, ensure_ascii=False, indent=2), encoding="utf-8")

    script_path = ROOT_DIR / "export_ppt_report.mjs"
    result = subprocess.run(
        [str(node_exe), str(script_path), str(spec_path)],
        cwd=str(ROOT_DIR),
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or "PPT 导出失败")
    return output_dir / "chart_report.pptx"


def build_download_url(path: Path) -> str:
    return f"/downloads/{quote(path.relative_to(ROOT_DIR).as_posix())}"


def generate_exports(dataset_id: str, config: dict[str, Any], formats: list[str] | None = None) -> dict[str, str]:
    requested = set(formats or ["png", "excel", "pptx"])
    requested = requested & {"png", "excel", "pptx"}
    if not requested:
        requested = {"png", "excel", "pptx"}

    table, dataset = load_dataset(dataset_id)
    output_dir = EXPORT_DIR / dataset_id
    output_dir.mkdir(parents=True, exist_ok=True)
    title = config.get("title") or "多元化图表汇报"

    exports: dict[str, str] = {}
    image_path = output_dir / "chart_report.png"
    if requested & {"png", "pptx"}:
        image_path = export_png(table, config, output_dir, title)
        exports["pngUrl"] = build_download_url(image_path)

    if "excel" in requested:
        excel_path = export_excel(table, dataset, config, output_dir)
        exports["excelUrl"] = build_download_url(excel_path)

    if "pptx" in requested:
        if not image_path.exists():
            image_path = export_png(table, config, output_dir, title)
            exports["pngUrl"] = build_download_url(image_path)
        pptx_path = export_ppt(table, dataset, image_path, config, output_dir)
        exports["pptxUrl"] = build_download_url(pptx_path)

    return exports


class PlatformHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(WEB_DIR), **kwargs)

    def log_message(self, fmt: str, *args) -> None:
        sys.stdout.write("%s - - [%s] %s\n" % (self.client_address[0], self.log_date_time_string(), fmt % args))

    def send_json(self, payload: dict[str, Any], status_code: int = 200) -> None:
        body, status_code = json_response(payload, status_code)
        self.send_response(status_code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self) -> None:
        path = urlparse(self.path).path
        if path == "/api/health":
            self.send_json({"ok": True, "message": "平台运行正常", "ai": get_ai_status()})
            return
        if path.startswith("/downloads/"):
            self.serve_download(unquote(path[len("/downloads/"):] if path.startswith("/downloads/") else path))
            return
        if path in {"/", "/index.html"}:
            self.path = "/index.html"
        return super().do_GET()

    def serve_download(self, relative_path: str) -> None:
        file_path = (ROOT_DIR / relative_path).resolve()
        root_path = ROOT_DIR.resolve()
        if root_path not in file_path.parents or not file_path.exists() or not file_path.is_file():
            self.send_json({"ok": False, "error": "文件不存在。"}, 404)
            return

        content = file_path.read_bytes()
        content_type = "application/octet-stream"
        if file_path.suffix.lower() == ".png":
            content_type = "image/png"
        elif file_path.suffix.lower() == ".xlsx":
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif file_path.suffix.lower() == ".pptx":
            content_type = "application/vnd.openxmlformats-officedocument.presentationml.presentation"

        safe_filename = f"download{file_path.suffix.lower()}"
        encoded_filename = quote(file_path.name)
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Disposition", f"attachment; filename=\"{safe_filename}\"; filename*=UTF-8''{encoded_filename}")
        self.send_header("Content-Length", str(len(content)))
        self.end_headers()
        self.wfile.write(content)

    def do_POST(self) -> None:
        path = urlparse(self.path).path
        if path == "/api/ingest":
            self.handle_ingest()
            return
        if path == "/api/export":
            self.handle_export()
            return
        self.send_json({"ok": False, "error": "接口不存在。"}, 404)

    def handle_ingest(self) -> None:
        try:
            content_type = self.headers.get("Content-Type", "")
            content_length = int(self.headers.get("Content-Length", "0"))
            request_body = self.rfile.read(content_length)
            extraction_meta: dict[str, Any] = {"mode": "rule", "used": False}

            if content_type.startswith("multipart/form-data"):
                filename, file_bytes = parse_upload(content_type, request_body)
                save_path = UPLOAD_DIR / filename
                save_path.write_bytes(file_bytes)
                table, source_name, raw_text = file_to_table(save_path)
                if raw_text.strip():
                    ai_table, extraction_meta = maybe_extract_with_ai(raw_text, source_name)
                    if ai_table is not None:
                        table = ai_table
                    else:
                        extraction_meta["parser"] = "rule"
                source_kind = "file"
            else:
                payload = json.loads(request_body.decode("utf-8")) if request_body else {}
                text = (payload.get("text") or "").strip()
                mode = str(payload.get("mode") or "auto").strip().lower()
                table, extraction_meta = parse_text_input_by_mode(text, "手动输入文本", mode)
                source_name = "手动输入文本"
                raw_text = text
                source_kind = "text"

            dataset_id = uuid.uuid4().hex[:12]
            dataset = build_dataset(
                table,
                source_name,
                raw_text,
                dataset_id=dataset_id,
                source_kind=source_kind,
                extraction_meta=extraction_meta,
            )
            save_dataset(dataset_id, clean_table(table), dataset)
            self.send_json({"ok": True, "dataset": dataset})
        except Exception as exc:
            self.send_json({"ok": False, "error": str(exc)}, 400)

    def handle_export(self) -> None:
        try:
            content_length = int(self.headers.get("Content-Length", "0"))
            payload = json.loads(self.rfile.read(content_length).decode("utf-8")) if content_length else {}
            dataset_id = payload.get("datasetId")
            if not dataset_id:
                raise ValueError("请先上传或解析一份数据。")
            config = {
                "type": payload.get("chartType") or "bar",
                "xField": payload.get("xField") or "",
                "yFields": payload.get("yFields") or [],
                "title": payload.get("title") or "多元化图表汇报",
            }
            formats = payload.get("formats")
            if formats is not None and not isinstance(formats, list):
                raise ValueError("导出格式参数不正确。")
            self.send_json({"ok": True, "exports": generate_exports(dataset_id, config, formats)})
        except Exception as exc:
            self.send_json({"ok": False, "error": str(exc)}, 400)


def main() -> None:
    ensure_dirs()
    server = ThreadingHTTPServer((HOST, PORT), PlatformHandler)
    print(f"数据智能分析平台已启动：http://{HOST}:{PORT}")
    server.serve_forever()


if __name__ == "__main__":
    main()
